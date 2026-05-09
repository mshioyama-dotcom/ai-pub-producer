# -*- coding: utf-8 -*-
"""
Markdown → DOCX 変換スクリプト
設計書（LifeBookNavigator_全体設計書_2.md）を Word ファイルに変換する。

サポートする Markdown 要素：
- 見出し（# / ## / ### / #### ）
- 表（パイプ区切り）
- コードブロック（```）
- 順序なしリスト（- / *）
- 順序付きリスト（1. 2. 3.）
- 引用ブロック（> ）
- インライン強調：**bold**, *italic*, ~~strike~~, `code`
- 区切り線（---）
- 簡易目次（先頭に挿入）
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor


# ---------- スタイル設定 ----------

def setup_styles(doc):
    """日本語向けに見栄えするスタイル設定"""
    styles = doc.styles

    # 既定本文
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal.font.size = Pt(10.5)
    # 日本語フォント設定（東アジア用）
    rPr = normal.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), "Yu Gothic")
    rFonts.set(qn("w:hAnsi"), "Yu Gothic")
    rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    rFonts.set(qn("w:cs"), "Yu Gothic")

    # 見出しスタイルのフォントを統一
    for h_name in ["Heading 1", "Heading 2", "Heading 3", "Heading 4"]:
        try:
            s = styles[h_name]
            s.font.name = "Yu Gothic"
            rPr = s.element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:ascii"), "Yu Gothic")
            rFonts.set(qn("w:hAnsi"), "Yu Gothic")
            rFonts.set(qn("w:eastAsia"), "Yu Gothic")
            rFonts.set(qn("w:cs"), "Yu Gothic")
            # 見出しは黒で読みやすく
            s.font.color.rgb = RGBColor(0x1F, 0x36, 0x5F)
            s.font.bold = True
        except KeyError:
            pass


# ---------- インライン記法のパース ----------

INLINE_PATTERN = re.compile(
    r"(\*\*([^*]+)\*\*|"          # **bold**
    r"~~([^~]+)~~|"                # ~~strike~~
    r"\*([^*]+)\*|"                # *italic*
    r"`([^`]+)`)"                  # `code`
)


def add_runs(paragraph, text):
    """インライン記法をパースして runs を追加"""
    if not text:
        return
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        # 前のプレーン部分
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
        # マッチ部分
        if m.group(2) is not None:  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3) is not None:  # strike
            run = paragraph.add_run(m.group(3))
            run.font.strike = True
        elif m.group(4) is not None:  # italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(5) is not None:  # inline code
            run = paragraph.add_run(m.group(5))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ---------- 各種ブロックの追加 ----------

def add_heading(doc, text, level):
    """見出しを追加（インライン記法対応）"""
    h = doc.add_heading(level=min(level, 4))
    add_runs(h, text)
    return h


def add_code_block(doc, lines, lang=None):
    """コードブロックを追加（等幅フォント・薄い背景）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    # 段落の背景を薄いグレーに
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_table(doc, rows):
    """Markdown 表を Word の表に変換（罫線・ヘッダー強調）"""
    if not rows:
        return
    # ヘッダー行（最初）と区切り行（2行目）を分離
    header = rows[0]
    body = rows[2:] if len(rows) > 2 else []
    cols = len(header)
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.style = "Light Grid Accent 1"
    # ヘッダー
    for i, cell_text in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        add_runs(p, cell_text.strip())
        for r in p.runs:
            r.bold = True
    # ボディ
    for row_idx, row in enumerate(body, start=1):
        for col_idx in range(cols):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            text = row[col_idx].strip() if col_idx < len(row) else ""
            p = cell.paragraphs[0]
            add_runs(p, text)
    doc.add_paragraph()  # 表の後の空行


def add_quote(doc, text):
    """引用ブロック"""
    p = doc.add_paragraph()
    try:
        p.style = doc.styles["Intense Quote"]
    except KeyError:
        try:
            p.style = doc.styles["Quote"]
        except KeyError:
            pass
    add_runs(p, text)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    add_runs(p, text)


def add_number(doc, text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    add_runs(p, text)


def add_separator(doc):
    """水平線"""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_paragraph(doc, text):
    """通常段落"""
    if not text.strip():
        return
    p = doc.add_paragraph()
    add_runs(p, text)


# ---------- 目次（TOC） ----------

def extract_headings(lines):
    """全見出しを抽出（コードブロック内は除く）"""
    in_code = False
    out = []
    for line in lines:
        s = line.rstrip("\n")
        if s.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            continue
        m = re.match(r"^(#{1,4})\s+(.+)$", s)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # インライン記法を除去
            text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
            text = re.sub(r"~~([^~]+)~~", r"\1", text)
            text = re.sub(r"\*([^*]+)\*", r"\1", text)
            text = re.sub(r"`([^`]+)`", r"\1", text)
            out.append((level, text))
    return out


def add_toc(doc, headings):
    """簡易目次（見出し一覧）を追加"""
    doc.add_heading("目次", level=1)
    for level, text in headings:
        if level > 3:  # 第4階層以上は目次から除外
            continue
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm((level - 1) * 0.6)
        run = p.add_run(text)
        if level == 1:
            run.bold = True
            run.font.size = Pt(11)
        elif level == 2:
            run.font.size = Pt(10.5)
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# ---------- メイン変換ロジック ----------

TABLE_LINE = re.compile(r"^\s*\|.*\|\s*$")
TABLE_DIVIDER = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")
HEADING = re.compile(r"^(#{1,6})\s+(.+)$")
BULLET = re.compile(r"^(\s*)[-*]\s+(.+)$")
NUMBER = re.compile(r"^(\s*)\d+\.\s+(.+)$")
QUOTE = re.compile(r"^>\s?(.*)$")
SEPARATOR = re.compile(r"^[-*_]{3,}\s*$")
CODE_FENCE = re.compile(r"^```(\w*)\s*$")


def convert(md_path, docx_path):
    md_text = Path(md_path).read_text(encoding="utf-8")
    lines = md_text.split("\n")

    doc = Document()
    setup_styles(doc)

    # 表紙タイトル
    title = doc.add_heading("AI出版プロデューサー / Life Book Navigator", level=0)
    sub = doc.add_paragraph()
    sub_run = sub.add_run("全体設計書 v2")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()  # 空行
    doc.add_paragraph(f"出力: docs/LifeBookNavigator_全体設計書_2.docx").runs[0].font.size = Pt(9)

    # 目次
    headings = extract_headings(lines)
    add_toc(doc, headings)

    # 改ページ
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # 状態管理
    in_code = False
    code_lines = []
    code_lang = None

    table_buf = []

    quote_buf = []

    def flush_table():
        nonlocal table_buf
        if table_buf:
            # 各行を | で分割
            parsed_rows = []
            for row in table_buf:
                # 前後の | を除去して分割
                cells = re.split(r"\s*\|\s*", row.strip().strip("|"))
                parsed_rows.append(cells)
            add_table(doc, parsed_rows)
            table_buf = []

    def flush_quote():
        nonlocal quote_buf
        if quote_buf:
            add_quote(doc, "\n".join(quote_buf))
            quote_buf = []

    for line in lines:
        s = line.rstrip("\n")

        # コードブロック内の処理
        if in_code:
            if CODE_FENCE.match(s):
                add_code_block(doc, code_lines, code_lang)
                code_lines = []
                code_lang = None
                in_code = False
            else:
                code_lines.append(line.rstrip("\n"))
            continue

        # コードブロック開始
        m = CODE_FENCE.match(s)
        if m:
            flush_table()
            flush_quote()
            in_code = True
            code_lang = m.group(1) or None
            continue

        # 表（連続する | 行）
        if TABLE_LINE.match(s):
            flush_quote()
            table_buf.append(s)
            continue
        else:
            flush_table()

        # 引用
        m = QUOTE.match(s)
        if m:
            quote_buf.append(m.group(1))
            continue
        else:
            flush_quote()

        # 区切り線
        if SEPARATOR.match(s):
            add_separator(doc)
            continue

        # 見出し
        m = HEADING.match(s)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2).strip(), level)
            continue

        # 順序なしリスト
        m = BULLET.match(s)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            add_bullet(doc, m.group(2), level=level)
            continue

        # 順序付きリスト
        m = NUMBER.match(s)
        if m:
            indent = len(m.group(1))
            level = indent // 2
            add_number(doc, m.group(2), level=level)
            continue

        # 通常段落
        if s.strip():
            add_paragraph(doc, s)
        else:
            # 空行は段落の区切りとしてスキップ（python-docx は自動で間隔を取る）
            pass

    # 最後のフラッシュ
    flush_table()
    flush_quote()

    doc.save(docx_path)
    print(f"出力完了: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
